"""
Inject Anggaran table into section F
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

DOCX_PATH = r"E:\UNIKOM\semester 6\Manajemen Proyek Perangkat Lunak\TUBES\SPMP - Sistem Informasi Manajemen Kost.docx"
doc = Document(DOCX_PATH)
body = doc.element.body
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def escape_xml(s):
    return s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace('"',"&quot;")

def make_p(text, style=None, size=11, bold=False, space_after=120):
    p_xml = f'<w:p {nsdecls("w")}><w:pPr>'
    if style:
        p_xml += f'<w:pStyle w:val="{style}"/>'
    p_xml += f'<w:spacing w:after="{space_after}"/>'
    p_xml += '</w:pPr>'
    p_xml += f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
    p_xml += f'<w:sz w:val="{size*2}"/>'
    if bold:
        p_xml += '<w:b/>'
    p_xml += '</w:rPr>'
    p_xml += f'<w:t xml:space="preserve">{escape_xml(text)}</w:t>'
    p_xml += '</w:r></w:p>'
    return parse_xml(p_xml)

def make_p_mixed(parts, space_after=120):
    runs_xml = ""
    for text, bold, size in parts:
        runs_xml += f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        runs_xml += f'<w:sz w:val="{size*2}"/>'
        if bold:
            runs_xml += '<w:b/>'
        runs_xml += '</w:rPr>'
        runs_xml += f'<w:t xml:space="preserve">{escape_xml(text)}</w:t></w:r>'
    p_xml = f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="{space_after}"/></w:pPr>{runs_xml}</w:p>'
    return parse_xml(p_xml)

def make_table_html(headers, rows, col_widths_twips=None, header_shade="D9E2F3"):
    """Create table. rows = [(values_list, shade_or_None), ...]"""
    ncols = len(headers)
    tbl_xml = f'<w:tbl {nsdecls("w")}>'
    tbl_xml += '<w:tblPr>'
    tbl_xml += '<w:tblStyle w:val="TableGrid"/>'
    tbl_xml += '<w:tblW w:w="5000" w:type="pct"/>'
    tbl_xml += '<w:jc w:val="center"/>'
    tbl_xml += '</w:tblPr><w:tblGrid>'
    for w in col_widths_twips or [1000]*ncols:
        tbl_xml += f'<w:gridCol w:w="{w}"/>'
    tbl_xml += '</w:tblGrid>'

    # Header row
    tbl_xml += '<w:tr>'
    for h in headers:
        tbl_xml += '<w:tc><w:tcPr><w:tcBorders>'
        for s in ['top','left','bottom','right']:
            tbl_xml += f'<w:{s} w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        tbl_xml += '</w:tcBorders><w:vAlign w:val="center"/>'
        tbl_xml += f'<w:shd w:fill="{header_shade}" w:val="clear"/>'
        tbl_xml += '</w:tcPr>'
        tbl_xml += '<w:p><w:pPr><w:jc w:val="center"/></w:pPr>'
        tbl_xml += '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        tbl_xml += '<w:sz w:val="16"/><w:b/></w:rPr>'
        tbl_xml += f'<w:t>{escape_xml(h)}</w:t></w:r></w:p></w:tc>'
    tbl_xml += '</w:tr>'

    # Data rows
    for r_data, row_shade in rows:
        tbl_xml += '<w:tr>'
        for ci, val in enumerate(r_data):
            tbl_xml += '<w:tc><w:tcPr><w:tcBorders>'
            for s in ['top','left','bottom','right']: 
                tbl_xml += f'<w:{s} w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            tbl_xml += '</w:tcBorders><w:vAlign w:val="center"/>'
            if row_shade:
                tbl_xml += f'<w:shd w:fill="{row_shade}" w:val="clear"/>'
            tbl_xml += '</w:tcPr>'
            hal = 'left' if ci in [1, 5] else 'right' if ci in [2, 3, 4] else 'center'
            tbl_xml += f'<w:p><w:pPr><w:jc w:val="{hal}"/></w:pPr>'
            tbl_xml += '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            tbl_xml += '<w:sz w:val="16"/></w:rPr>'
            tbl_xml += f'<w:t>{escape_xml(str(val))}</w:t></w:r></w:p></w:tc>'
        tbl_xml += '</w:tr>'
    tbl_xml += '</w:tbl>'
    return parse_xml(tbl_xml)

# ===== FIND F. RENCANA ANGGARAN =====
f_elem = None
placeholder_elem = None
for para in doc.paragraphs:
    if "F. RENCANA ANGGARAN" in para.text:
        f_elem = para._element
    if f_elem is not None and "[Akan dilengkapi kemudian.]" in para.text:
        placeholder_elem = para._element
        break

if f_elem is None:
    print("❌ Section F not found")
    exit(1)

# Remove the placeholder
if placeholder_elem is not None:
    try:
        body.remove(placeholder_elem)
        print("✅ Removed placeholder")
    except:
        pass

# ===== INSERT CONTENT BEFORE G. PENUTUP =====
# Find G. PENUTUP
g_elem = None
for para in doc.paragraphs:
    if "G. PENUTUP" in para.text:
        g_elem = para._element
        break

if g_elem is None:
    print("❌ Section G not found")
    exit(1)

# Narrative
narasi = make_p("Rencana anggaran proyek Sistem Informasi Manajemen Kost ini disusun berdasarkan estimasi biaya pengembangan selama 84 hari kerja (6 April – 30 Juli 2026) dengan melibatkan 5 anggota tim. Biaya Sumber Daya Manusia (SDM) dihitung berdasarkan Upah Minimum Regional (UMR) Kota Bandung tahun 2026 yang ditetapkan sebesar Rp 4.292.000/bulan [5]. Tarif per jam dihitung dengan asumsi 173 jam kerja per bulan sesuai ketentuan Undang-Undang Ketenagakerjaan.", None, 11, False)
g_elem.addprevious(narasi)

# UMR calc
# Rp 4.292.000 / 173 jam = ~Rp 24.809/jam → ~Rp 24.800/jam
# But let me use a more careful number
umr_per_bulan = 4292000
jam_per_bulan = 173
tarif_per_jam = int(round(umr_per_bulan / jam_per_bulan, -2))  # round to nearest 100

# Table 1: SDM
sdm_data = [
    ("Jeri Gunawan", "Fullstack Developer", "57", f"Rp {tarif_per_jam:,}", f"Rp {tarif_per_jam * 57:,}"),
    ("Raden Ris Ravanel", "UI/UX Designer", "15", f"Rp {tarif_per_jam:,}", f"Rp {tarif_per_jam * 15:,}"),
    ("Faliq Husnan", "Database Engineer", "12", f"Rp {tarif_per_jam:,}", f"Rp {tarif_per_jam * 12:,}"),
    ("Alghifari M. B. S.", "QA Engineer", "12", f"Rp {tarif_per_jam:,}", f"Rp {tarif_per_jam * 12:,}"),
    ("Herdiyan Aditya", "Project Manager", "9", f"Rp {tarif_per_jam:,}", f"Rp {tarif_per_jam * 9:,}"),
]
total_sdm = sum(tarif_per_jam * int(row[2]) for row in sdm_data)

sdm_rows2 = []
for i, (nama, peran, jam, tarif, total) in enumerate(sdm_data):
    sdm_rows2.append(([str(i+1), nama, peran, jam, tarif, total], None))
sdm_rows2.append((["", "Total Biaya SDM", "", "", "", f"Rp {total_sdm:,}"], "E2EFDA"))

sdm_tbl = make_table_html(
    ["No", "Nama", "Peran", "Jam", "Tarif/Jam", "Total"],
    sdm_rows2,
    [600, 2200, 2000, 800, 1500, 1500]
)
g_elem.addprevious(sdm_tbl)

# SDM heading
sdm_heading = make_p("Tabel Rencana Anggaran Biaya SDM", "Heading4", 11, True)
g_elem.addprevious(sdm_heading)

# Table 2: Non-SDM
non_sdm = [
    ("Domain (.com / .id)", "1 tahun", "Rp 250.000", "Rp 250.000"),
    ("VPS (1 bulan)", "3 bulan", "Rp 150.000/bln", "Rp 450.000"),
    ("SSL Certificate (Let's Encrypt)", "—", "Gratis", "Rp 0"),
    ("Tools & Software (Open Source)", "—", "Gratis", "Rp 0"),
    ("Biaya Operasional (listrik, internet)", "3 bulan", "Rp 200.000/bln", "Rp 600.000"),
    ("Cadangan (10%)", "—", "—", f"Rp {int((total_sdm + 250000 + 450000 + 600000) * 0.1):,}"),
]

non_sdm_heading = make_p("Tabel Rencana Anggaran Biaya Non-SDM", "Heading4", 11, True)
g_elem.addprevious(non_sdm_heading)

total_non = 250000 + 450000 + 0 + 0 + 600000
cadangan = int((total_sdm + total_non) * 0.1)
total_all = total_sdm + total_non + cadangan

non_sdm2 = [
    ("Domain (.com / .id)", "1 tahun", "Rp 250.000", "Rp 250.000"),
    ("VPS Hosting", "3 bulan", "Rp 150.000/bln", "Rp 450.000"),
    ("SSL Certificate", "—", "Gratis", "Rp 0"),
    ("Tools & Software", "—", "Open Source", "Rp 0"),
    ("Operasional", "3 bulan", "Rp 200.000/bln", "Rp 600.000"),
]

non_rows = []
for i, (item, qty, satuan, total) in enumerate(non_sdm2):
    non_rows.append(([str(i+1), item, qty, satuan, total], None))
non_rows.append((["", "Cadangan (10%)", "", "", f"Rp {cadangan:,}"], None))
non_rows.append((["", "Total Biaya Non-SDM", "", "", f"Rp {total_non + cadangan:,}"], "E2EFDA"))

non_tbl = make_table_html(
    ["No", "Item", "Kuantitas", "Satuan", "Total"],
    non_rows,
    [600, 2500, 1800, 2000, 1500]
)
g_elem.addprevious(non_tbl)

# Grand total
grand_p = make_p_mixed([
    ("Grand Total Anggaran: Rp ", False, 11),
    (f"{total_all:,}", True, 12),
    (" (termasuk cadangan 10%).", False, 11),
])
g_elem.addprevious(grand_p)

# Note about UMR
umr_note = make_p(f"Catatan: Tarif SDM dihitung berdasarkan UMR Kota Bandung 2026 sebesar Rp {umr_per_bulan:,}/bulan dengan 173 jam kerja per bulan, sehingga diperoleh tarif Rp {tarif_per_jam:,}/jam [5]. Biaya aktual dapat berbeda tergantung kebijakan perusahaan dan negosiasi kontrak.", None, 10, False)
g_elem.addprevious(umr_note)

doc.save(DOCX_PATH)
print(f"✅ Anggaran injected!")
print(f"   Tarif SDM: Rp {tarif_per_jam:,}/jam (UMR Bandung 2026: Rp {umr_per_bulan:,}/bln)")
print(f"   Total SDM: Rp {total_sdm:,}")
print(f"   Total Non-SDM + Cadangan: Rp {total_non + cadangan:,}")
print(f"   Grand Total: Rp {total_all:,}")
