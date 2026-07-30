"""
Create: Anggaran dan Monitoring - Sistem Informasi Manajemen Kost.docx
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
from collections import OrderedDict

SRC = r"E:\UNIKOM\semester 6\Manajemen Proyek Perangkat Lunak\TUBES\SPMP - Sistem Informasi Manajemen Kost.docx"
OUT = r"E:\UNIKOM\semester 6\Manajemen Proyek Perangkat Lunak\TUBES\Anggaran dan Monitoring - Sistem Informasi Manajemen Kost.docx"

src_doc = Document(SRC)
doc = Document()

# ─── EXTRACT WBS DATA ───
wbs_tasks = []
for table in src_doc.tables:
    first_row = "|".join(c.text.strip() for c in table.rows[0].cells) if table.rows else ""
    if "Rincian" in first_row or ("Kode" in first_row and "Nama Aktivitas" in first_row):
        for i, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) < 6:
                continue
            kode = cells[1].text.strip()
            nama = cells[2].text.strip()
            pic = cells[3].text.strip()
            dur_text = cells[5].text.strip()
            if not nama or nama in ["Nama Aktivitas", ""]:
                continue
            if kode in ["Kode", "1.0", "2.0", "3.0", "4.0", "5.0", ""]:
                continue
            d_match = re.search(r'(\d+)', dur_text)
            dur = int(d_match.group(1)) if d_match else 0
            if kode and nama and pic and dur > 0:
                wbs_tasks.append((kode, nama, pic, dur))
        break

# ─── HELPERS ───
def add_h(text, level=1):
    doc.add_heading(text, level=level)

def add_p(text, size=11, bold=False, italic=False, space_after=120):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after/20)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_tbl(headers, rows, col_cm=None):
    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)
        run.font.bold = True
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, (values, shade) in enumerate(rows, 1):
        for ci, val in enumerate(values):
            cell = table.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.alignment = align
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(7)
            if shade:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade}" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_cm:
        for row in table.rows:
            for i, w in enumerate(col_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table

def rp(val):
    s = f"{int(val):,}".replace(",", ".")
    return f"Rp {s}"

tarif = int(round(4_292_000 / 173, -2))

# ═══════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAMPIRAN ANGGARAN & MONITORING")
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.font.bold = True
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistem Informasi Manajemen Kost Berbasis Web")
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.italic = True
doc.add_paragraph()

# ─── A. PARAMETER ───
add_h("A. Parameter Dasar", 1)
add_p("UMR Kota Bandung 2026: Rp 4.292.000/bulan", 10)
add_p("Jam kerja per bulan: 173 jam", 10)
add_p(f"Tarif per jam: Rp {24_800:,} (dibulatkan ke ratusan terdekat)", 10, bold=True)
add_p(f"Total jam proyek: {sum(t[3] for t in wbs_tasks)} jam ({len(wbs_tasks)} task, 5 fase)", 10)
add_p("Durasi proyek: 84 hari kerja (6 Apr – 30 Jul 2026)", 10)
add_p("", 6)

# ─── B. ANGGARAN PER TASK ───
add_h("B. Rincian Anggaran per Task", 1)

task_rows = []
phases = OrderedDict()
for kode, nama, pic, dur in wbs_tasks:
    biaya = dur * tarif
    pnum = kode.split(".")[0]
    pmap = {"1": "Fase 1", "2": "Fase 2", "3": "Fase 3", "4": "Fase 4", "5": "Fase 5"}
    ph = pmap.get(pnum, "?")
    if ph not in phases:
        phases[ph] = {"jam": 0, "biaya": 0}
    phases[ph]["jam"] += dur
    phases[ph]["biaya"] += biaya
    shade = "FFF3E0" if biaya >= 100_000 else None
    task_rows.append(([kode, nama[:55], pic, str(dur), rp(tarif), rp(biaya)], shade))

add_tbl(["Kode", "Nama Aktivitas", "PIC", "Jam", "Tarif", "Biaya"], task_rows, [1.5, 7.5, 2.0, 1.0, 2.0, 2.5])
add_p("", 6)

add_h("Rekapitulasi per Fase", 2)
phase_rows = []
sub_total = 0
for ph, d in phases.items():
    phase_rows.append(([ph, str(d["jam"]), rp(d["biaya"])], "E8F5E9"))
    sub_total += d["biaya"]
add_tbl(["Fase", "Jam", "Biaya"], phase_rows, [5.0, 3.0, 5.0])
add_p("", 6)

add_p(f"Total Biaya SDM: {rp(sub_total)}", 11, bold=True)
add_p(f"Cadangan (10%): {rp(int(sub_total*0.1))}", 10)
add_p(f"Grand Total Anggaran: {rp(int(sub_total*1.1))}", 12, bold=True)
add_p("", 6)

# ─── C. EVM ───
add_h("C. Rencana Monitoring & Earned Value Management", 1)
add_p("Earned Value Management (EVM) digunakan untuk memantau kemajuan proyek secara objektif dengan membandingkan biaya dan jadwal yang direncanakan dengan aktual pada setiap milestone.", 10)

add_h("C.1 Metrik EVM", 2)
metrics = [
    ("PV (Planned Value)", "Nilai anggaran yang seharusnya sudah terselesaikan sesuai jadwal."),
    ("EV (Earned Value)", "Nilai pekerjaan yang benar-benar telah selesai (% progres × total anggaran fase)."),
    ("AC (Actual Cost)", "Biaya aktual yang telah dikeluarkan."),
    ("SV = EV − PV", "SV > 0 = lebih cepat; SV < 0 = terlambat."),
    ("CV = EV − AC", "CV > 0 = di bawah budget; CV < 0 = over budget."),
    ("SPI = EV / PV", "SPI > 1 = lebih cepat; SPI < 1 = terlambat. Toleransi ≥ 0,9."),
    ("CPI = EV / AC", "CPI > 1 = efisien; CPI < 1 = boros. Toleransi ≥ 0,9."),
]
for m, d in metrics:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"• {m}: ")
    r.font.bold = True; r.font.size = Pt(10)
    r = p.add_run(d)
    r.font.size = Pt(10)

# ─── D. PV BASELINE ───
add_h("C.2 Planned Value Baseline per Fase", 2)

fase_schedule = [
    ("Fase 1: Perencanaan", 0, 9, phases.get("Fase 1", {"jam":0,"biaya":0})),
    ("Fase 2: Analisis & Desain", 9, 15, phases.get("Fase 2", {"jam":0,"biaya":0})),
    ("Fase 3: Implementasi", 24, 46-24, phases.get("Fase 3", {"jam":0,"biaya":0})),
    ("Fase 4: Pengujian", 69, 81-69, phases.get("Fase 4", {"jam":0,"biaya":0})),
    ("Fase 5: Deployment", 81, 84-81, phases.get("Fase 5", {"jam":0,"biaya":0})),
]

pv_rows = []
cum_pv = 0
for name, st, dur, d in fase_schedule:
    cum_pv += d["biaya"]
    pv_rows.append(([name, f"H-{st}", f"H-{st+dur}", str(dur), str(d["jam"]), rp(d["biaya"]), rp(cum_pv)], "E3F2FD"))

add_tbl(["Fase", "Mulai", "Selesai", "Dur (hr)", "Jam", "Biaya", "PV Kum."], pv_rows, [4.0, 1.5, 1.5, 1.5, 1.5, 2.5, 2.5])
add_p(f"Total PV Baseline: {rp(cum_pv)}", 10, bold=True)
add_p("", 6)

# ─── E. MILESTONES ───
add_h("C.3 Jadwal Monitoring (Milestone EVM)", 2)

milestones = [
    ("M1", "Akhir Fase 1", "H-9", "16 Apr", "Herdiyan, Semua", str(phases.get("Fase 1",{"jam":0})["jam"]), phases.get("Fase 1",{"biaya":0})["biaya"]),
    ("M2", "Akhir Fase 2", "H-24", "7 Mei", "Raden, Faliq, Herdiyan", str(phases.get("Fase 2",{"jam":0})["jam"]), phases.get("Fase 2",{"biaya":0})["biaya"]),
    ("M3", "Akhir Fase 3", "H-69", "9 Jul", "Jeri", str(phases.get("Fase 3",{"jam":0})["jam"]), phases.get("Fase 3",{"biaya":0})["biaya"]),
    ("M4", "Akhir Fase 4", "H-81", "23 Jul", "Alghifari, Semua", str(phases.get("Fase 4",{"jam":0})["jam"]), phases.get("Fase 4",{"biaya":0})["biaya"]),
    ("M5", "Akhir Fase 5", "H-84", "30 Jul", "Semua", str(phases.get("Fase 5",{"jam":0})["jam"]), phases.get("Fase 5",{"biaya":0})["biaya"]),
]

mln_rows = []
c = 0
for m in milestones:
    c += m[6]
    mln_rows.append(([m[0], m[1], m[2], m[3], m[4], m[5], rp(m[6]), rp(c)], None))

add_tbl(["MS", "Deskripsi", "Hari", "Tgl", "PIC", "Jam", "Biaya", "PV Kum."], mln_rows, [1.5, 3.5, 1.5, 2.0, 3.0, 1.5, 2.5, 2.5])
add_p("", 6)

# ─── F. TEMPLATE ───
add_h("C.4 Template Laporan Monitoring per Milestone", 2)
add_p("Setiap akhir fase, tim mengisi tabel EVM berikut:", 10)

trows = [
    (["PV (Planned Value)", "(dari baseline)", "Anggaran yang seharusnya sudah dipakai"], None),
    (["EV (Earned Value)", "(hitung dari % progres)", "Nilai pekerjaan yang sudah selesai"], None),
    (["AC (Actual Cost)", "(dari pengeluaran riil)", "Biaya yang sudah dikeluarkan"], None),
    (["SV = EV − PV", "", "> 0 = lebih cepat < 0 = terlambat"], "FFF9C4"),
    (["CV = EV − AC", "", "> 0 = di bawah budget < 0 = over budget"], "FFF9C4"),
    (["SPI = EV / PV", "", "> 1 = lebih cepat < 1 = terlambat"], "FFF9C4"),
    (["CPI = EV / AC", "", "> 1 = efisien < 1 = boros"], "FFF9C4"),
    (["Tindakan Koreksi", "", "Jika SPI/CPI < 0,9 → realokasi/jadwal ulang"], "FFEBEE"),
]
add_tbl(["Metrik", "Cara Hitung", "Interpretasi"], trows, [4.0, 4.0, 6.0])

add_p("Toleransi SPI dan CPI ≥ 0,9. Jika di bawah, perlu tindakan koreksi segera.", 9, italic=True)
doc.add_paragraph()
add_p("— Dokumen ini melengkapi SPMP —", 10, italic=True)

# ─── SAVE ───
doc.save(OUT)
print(f"✅ Saved: {OUT}")
print(f"   {len(wbs_tasks)} tasks with budget per task")
print(f"   EVM baseline + monitoring template")
