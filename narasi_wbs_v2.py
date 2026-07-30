"""
Replace intro text at C.3 with proper narrative before WBS table
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time

PATH = r"E:\UNIKOM\semester 6\Manajemen Proyek Perangkat Lunak\TUBES\SPMP - Sistem Informasi Manajemen Kost.docx"

time.sleep(1)  # give any lock a moment

doc = Document(PATH)

def make_para(text, bold=False, italic=False, size=20, spacing_after=120):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if spacing_after:
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:after'), str(spacing_after))
        pPr.append(sp)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def make_mixed_para(segments, spacing_after=120):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if spacing_after:
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:after'), str(spacing_after))
        pPr.append(sp)
    p.append(pPr)
    for text, bold, italic in segments:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            rPr.append(OxmlElement('w:b'))
        if italic:
            rPr.append(OxmlElement('w:i'))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    return p

# ─── Find C.3 heading ───
c3_idx = None
for i, p in enumerate(doc.paragraphs):
    if 'C.3' in p.text and 'WBS' in p.text:
        c3_idx = i
        break

if c3_idx is None:
    print("ERROR: C.3 heading not found")
    exit(1)

# Find D section
d_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('D.') and 'PENJADWALAN' in p.text.upper():
        d_idx = i
        break

c3_elem = doc.paragraphs[c3_idx]._element
d_elem = doc.paragraphs[d_idx]._element

# Remove everything between C.3 and D
body = doc.element.body
to_remove = []
found_c3 = False
for elem in list(body):
    if elem is c3_elem:
        found_c3 = True
        continue
    if elem is d_elem:
        break
    if found_c3:
        to_remove.append(elem)

print(f"Removing {len(to_remove)} elements between C.3 and D")
for elem in to_remove:
    body.remove(elem)

# ─── Insert new content ───
insert_after = c3_elem

# Narrative paragraph 1
n1 = make_mixed_para([
    ("Work Breakdown Structure (WBS) adalah ", False, False),
    ("pembagian pekerjaan proyek secara hierarkis ", False, True),
    ("dari level tertinggi (proyek) hingga level terendah (work package) [4]. ", False, False),
    ("WBS proyek Sistem Informasi Manajemen Kost ini terdiri dari ", False, False),
    ("5 fase utama, 66 task", True, False),
    (", dengan total ", False, False),
    ("103 jam kerja", True, False),
    (" yang dikerjakan oleh tim beranggotakan ", False, False),
    ("5 orang", True, False),
    (" — Herdiyan (PM/Ketua), Jeri (Fullstack), Alghifari (QA), Raden (UI/UX), dan Faliq (Database).", False, False),
])
insert_after.addnext(n1)
insert_after = n1

# Narrative paragraph 2
n2 = make_mixed_para([
    ("Kelima fase tersebut meliputi: (1) Perencanaan", True, False),
    (" — inisiasi proyek, analisis kebutuhan, penyusunan SPMP; ", False, False),
    ("(2) Perancangan Sistem", True, False),
    (" — desain antarmuka admin & pengguna oleh Raden, serta persiapan lingkungan pengembangan oleh Jeri; ", False, False),
    ("(3) Pengembangan Sistem", True, False),
    (" — implementasi frontend dan backend oleh Jeri, dan basis data oleh Faliq secara paralel; ", False, False),
    ("(4) Pengujian Sistem dengan UAT", True, False),
    (" — unit testing, integration testing, dan user acceptance test oleh Alghifari; serta ", False, False),
    ("(5) Implementasi Sistem", True, False),
    (" — deployment server, dokumentasi, dan pelatihan penggunaan sistem.", False, False),
])
insert_after.addnext(n2)
insert_after = n2

# Narrative paragraph 3
n3 = make_mixed_para([
    ("Penjadwalan proyek menggunakan prinsip ", False, False),
    ("fase sekuensial", True, False),
    (" di mana fase berikutnya dimulai setelah fase sebelumnya selesai. ", False, False),
    ("Di dalam satu fase, task dengan penanggung jawab (PIC) yang berbeda dijalankan secara ", False, False),
    ("paralel", True, False),
    (" untuk mengoptimalkan waktu pengerjaan. ", False, False),
    ("Weekend (Sabtu-Minggu) tidak diperhitungkan sebagai hari kerja. ", False, False),
    ("Proyek dimulai pada ", False, False),
    ("Senin, 6 April 2026", True, False),
    (" dan ditargetkan selesai pada ", False, False),
    ("Kamis, 30 Juli 2026", True, False),
    (", dengan total ", False, False),
    ("17 minggu", True, False),
    (" (83 hari kerja efektif). ", False, False),
    ("Durasi setiap task dinyatakan dalam jam kerja dan dikonversi ke Minggu ke-", False, False),
    (" berdasarkan Minggu 1 = 6–10 April 2026.", False, True),
])
insert_after.addnext(n3)
insert_after = n3

# Spacer
sp = OxmlElement('w:p')
insert_after.addnext(sp)
insert_after = sp

# Subtitle
sub = make_para("Rincian Work Breakdown Structure", bold=True, size=22, spacing_after=60)
insert_after.addnext(sub)
insert_after = sub

# Info bar
info = make_para("Durasi proyek: 17 minggu  |  Platform: Web App Fullstack  |  Tim: 5 orang  |  Total: 66 task / 103 jam kerja", italic=True, size=18, spacing_after=80)
insert_after.addnext(info)

# Now re-create the table (same as before but rebuild it)
def make_cell(text, bold=False, center=False, size=16, fill=None, color=None):
    tc = OxmlElement('w:tc')
    tcPr = OxmlElement('w:tcPr')
    if fill:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    tc.append(tcPr)
    p = OxmlElement('w:p')
    if center:
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if color:
        c_e = OxmlElement('w:color')
        c_e.set(qn('w:val'), color)
        rPr.append(c_e)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = str(text)
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc

def create_table_element():
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'Table Grid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    tblBorders = OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblLook = OxmlElement('w:tblLook')
    tblLook.set(qn('w:val'), '04A0')
    tblPr.append(tblLook)
    tbl.append(tblPr)
    tblGrid = OxmlElement('w:tblGrid')
    for w in ['700','1200','6500','1600','1300','900','3000']:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), w)
        tblGrid.append(gc)
    tbl.append(tblGrid)
    return tbl

from datetime import date
START = date(2026, 4, 6)
def week_no(d_str):
    day_map = {'Jan':1,'Feb':2,'Mar':3,'Apr':4,'May':5,'Jun':6,'Jul':7,'Aug':8,'Sep':9,'Oct':10,'Nov':11,'Dec':12}
    parts = d_str.split()
    dd = date(2026, day_map[parts[1]], int(parts[0]))
    return (dd - START).days // 7 + 1

def wk(s, e):
    ws, we = week_no(s), week_no(e)
    return str(ws) if ws == we else f"{ws}–{we}"

# Task data
T = {}
def add(id, kode, nama, pic, dur, m, s, ket):
    T[id] = {"kode":kode,"nama":nama,"pic":pic,"dur":dur,"mulai":m,"selesai":s,"ket":ket}

add(1,"1.1","Mempelajari proses bisnis kost secara manual","Herdiyan",1,"06 Apr","06 Apr","Observasi langsung/manual")
add(2,"1.2","Mengidentifikasi kebutuhan pengguna (admin & penghuni)","Herdiyan",2,"07 Apr","08 Apr","Wawancara pemilik kos")
add(3,"1.3","Menentukan fitur utama (sewa, bayar, data, aduan)","Herdiyan",1,"09 Apr","09 Apr","Prioritas fitur")
add(4,"1.4","Menyusun perencanaan waktu & pembagian tugas tim","Herdiyan",2,"10 Apr","13 Apr","Scrum planning")
add(5,"1.5","Mengidentifikasi risiko & solusi permasalahan","Herdiyan",1,"14 Apr","14 Apr","Risk register awal")
add(6,"1.6","AI untuk cari referensi & ide sistem","Semua",1,"06 Apr","06 Apr","ChatGPT, Claude, Exa")
add(7,"1.7","Pencatatan Dokumentasi SPMP","Herdiyan",2,"15 Apr","16 Apr","Sesuai template IEEE 1058")
add(8,"2.1.1","Merancang wireframe admin kost","Raden",1,"17 Apr","17 Apr","Figma")
add(9,"2.1.2","Mendesain Dashboard admin","Raden",1,"20 Apr","20 Apr","Figma")
add(10,"2.1.3","Mendesain Manajemen Kamar","Raden",1,"21 Apr","21 Apr","Figma")
add(11,"2.1.4","Mendesain Manajemen Booking","Raden",1,"22 Apr","22 Apr","Figma")
add(12,"2.1.5","Mendesain Manajemen Pembayaran","Raden",1,"23 Apr","23 Apr","Figma")
add(13,"2.1.6","Mendesain Check-Out & Perpanjangan","Raden",1,"24 Apr","24 Apr","Figma")
add(14,"2.1.7","Mendesain Aduan Kerusakan","Raden",1,"27 Apr","27 Apr","Figma")
add(15,"2.1.8","Meninjau ulang desain bersama tim","Raden",1,"28 Apr","28 Apr","Review internal")
add(16,"2.2.1","Merancang wireframe penghuni kost","Raden",1,"29 Apr","29 Apr","Figma")
add(17,"2.2.2","Mendesain Autentikasi","Raden",1,"30 Apr","30 Apr","Figma")
add(18,"2.2.3","Mendesain Daftar Kamar","Raden",1,"01 May","01 May","Figma")
add(19,"2.2.4","Mendesain Booking Kamar","Raden",1,"04 May","04 May","Figma")
add(20,"2.2.5","Mendesain Pembayaran","Raden",1,"05 May","05 May","Figma")
add(21,"2.2.6","Mendesain Pengaduan Kerusakan","Raden",1,"06 May","06 May","Figma")
add(22,"2.2.7","Meninjau ulang desain bersama tim","Raden",1,"07 May","07 May","Review internal")
add(23,"2.3.1","Mempersiapkan sumber daya akun","Jeri",1,"17 Apr","17 Apr","Akun GitHub, Figma")
add(24,"2.3.2","Mempersiapkan IDE, Git, Github","Jeri",1,"20 Apr","20 Apr","VS Code, Laragon")
add(25,"2.3.3","Membuat repository Github","Jeri",1,"21 Apr","21 Apr","Branch: main/dev")
add(26,"2.3.4","Melakukan clone ke perangkat","Jeri",1,"22 Apr","22 Apr","Git clone")
add(27,"3.1.1","Instalasi Laravel & sumber daya","Jeri",1,"08 May","08 May","Composer, npm")
add(28,"3.1.2","Upload sistem ke repository Github","Jeri",1,"11 May","11 May","Push awal")
add(29,"3.1.3","Membuat halaman autentikasi (FE)","Jeri",2,"12 May","13 May","Login, register, reset")
add(30,"3.1.4","Membuat halaman tampilan admin","Jeri",4,"14 May","19 May","Dashboard, kamar, booking")
add(31,"3.1.5","Membuat halaman tampilan pengguna","Jeri",3,"20 May","22 May","Katalog, profil")
add(32,"3.1.6","Membuat fitur admin sesuai desain","Jeri",4,"25 May","28 May","Sesuai Figma")
add(33,"3.1.7","Membuat fitur pengguna sesuai desain","Jeri",3,"29 May","02 Jun","Sesuai Figma")
add(34,"3.1.8","Menghubungkan FE-BE (integrasi)","Jeri",2,"03 Jun","04 Jun","API endpoint testing")
add(35,"3.1.9","Pencatatan dokumentasi DPPL","Jeri",2,"05 Jun","08 Jun","Frontend docs")
add(36,"3.2.1","Membuat autentikasi Laravel Breeze","Jeri",2,"09 Jun","10 Jun","Middleware auth")
add(37,"3.2.2","Mengatur hak akses sesuai role","Jeri",1,"11 Jun","11 Jun","Admin & user gates")
add(38,"3.2.3","Membuat statistik dashboard","Jeri",2,"12 Jun","15 Jun","Chart.js")
add(39,"3.2.4","Membuat Modul Kamar","Jeri",3,"16 Jun","18 Jun","CRUD, filter, gambar")
add(40,"3.2.5","Membuat Modul Booking","Jeri",3,"19 Jun","23 Jun","Booking & approval")
add(41,"3.2.6","Membuat Modul Pembayaran","Jeri",4,"24 Jun","29 Jun","Upload & verifikasi")
add(42,"3.2.7","Membuat Check-Out & Perpanjangan","Jeri",3,"30 Jun","02 Jul","Status kamar")
add(43,"3.2.8","Membuat Modul Aduan Kerusakan","Jeri",2,"03 Jul","06 Jul","Lapor & tindak lanjut")
add(44,"3.2.9","AI untuk debugging kode","Jeri",1,"07 Jul","07 Jul","Claude / ChatGPT")
add(45,"3.2.10","Pencatatan dokumentasi DPPL","Jeri",2,"08 Jul","09 Jul","Backend docs")
add(46,"3.3.1","Merancang ERD","Faliq",1,"08 May","08 May","draw.io")
add(47,"3.3.2","Membuat skema relasi","Faliq",1,"11 May","11 May","MySQL Workbench")
add(48,"3.3.3","Membuat Migration","Faliq",2,"12 May","13 May","Laravel migration")
add(49,"3.3.4","Membuat Relasi & Constraint","Faliq",1,"14 May","14 May","Foreign key, index")
add(50,"3.3.5","Membuat Seeder","Faliq",1,"15 May","15 May","Data dummy")
add(51,"3.3.6","Membuat Query untuk semua role","Faliq",2,"18 May","19 May","Eloquent ORM")
add(52,"3.3.7","Membuat tabel fasilitas barang","Faliq",1,"20 May","20 May","Inventaris kamar")
add(53,"3.3.8","Dokumentasi Basis Data","Faliq",1,"21 May","21 May","ERD & skema")
add(54,"3.3.9","Username & password admin","Faliq",1,"22 May","22 May","Seeder admin")
add(55,"3.3.10","Pencatatan dokumentasi DPPL","Faliq",1,"25 May","25 May","DB docs")
add(56,"4.1","Membuat perencanaan pengujian","Alghifari",1,"10 Jul","10 Jul","Test plan document")
add(57,"4.2","Unit testing (per Method)","Alghifari",3,"13 Jul","15 Jul","Pest/PHPUnit")
add(58,"4.3","Uji alur Booking sampai Check-Out","Alghifari",2,"16 Jul","17 Jul","Integration test")
add(59,"4.4","Uji Perpanjangan Kontrak","Alghifari",1,"20 Jul","20 Jul","Skenario khusus")
add(60,"4.5","Uji Aduan Kerusakan","Alghifari",1,"21 Jul","21 Jul","Skenario khusus")
add(61,"4.6","UAT 2 role: Pengguna & Admin","Alghifari",2,"22 Jul","23 Jul","Black-box testing")
add(62,"4.7","Pencatatan dokumentasi DPPL","Alghifari",2,"24 Jul","27 Jul","Test docs")
add(63,"5.1","Deployment ke VPS","Jeri",1,"28 Jul","28 Jul","VPS / hosting")
add(64,"5.2","Konfigurasi domain & SSL","Jeri",1,"29 Jul","29 Jul","Domain, HTTPS")
add(65,"5.3","Pelatihan penggunaan sistem","Semua",1,"28 Jul","28 Jul","Demo ke owner")
add(66,"5.4","Dokumentasi sistem & buku manual","Semua",2,"29 Jul","30 Jul","User manual")

# Phase structure
phase_structure = [
    ("1.0", "PERENCANAAN", [(None, None, list(range(1, 8)))]),
    ("2.0", "PERANCANGAN SISTEM", [
        ("2.1", "Desain Admin Kost", list(range(8, 16))),
        ("2.2", "Desain Pengguna Kost", list(range(16, 23))),
        ("2.3", "Persiapan Lingkungan Pengembangan", list(range(23, 27))),
    ]),
    ("3.0", "PENGEMBANGAN SISTEM", [
        ("3.1", "Frontend Developer", list(range(27, 36))),
        ("3.2", "Backend Developer", list(range(36, 46))),
        ("3.3", "Database Developer", list(range(46, 56))),
    ]),
    ("4.0", "PENGUJIAN SISTEM DENGAN UAT", [(None, None, list(range(56, 63)))]),
    ("5.0", "IMPLEMENTASI SISTEM", [(None, None, list(range(63, 67)))]),
]

# Build table
tbl = create_table_element()

# Header
thr = OxmlElement('w:tr')
for h in ['No','Kode','Nama Aktivitas','PIC','Minggu Ke-','Dur.','Keterangan']:
    thr.append(make_cell(h, bold=True, center=True, size=16, fill='2E4057', color='FFFFFF'))
tbl.append(thr)

seq = 0
row_idx = 0
for pk, pn, subs in phase_structure:
    # Phase header
    tr_ph = OxmlElement('w:tr')
    for val in ['', pk, pn, '', '', '', '']:
        idx_ = 0
    tr_ph.append(make_cell('', fill='D9E2F3'))
    tr_ph.append(make_cell(pk, bold=True, fill='D9E2F3'))
    tr_ph.append(make_cell(pn, bold=True, fill='D9E2F3'))
    tr_ph.append(make_cell('', fill='D9E2F3'))
    tr_ph.append(make_cell('', fill='D9E2F3'))
    tr_ph.append(make_cell('', fill='D9E2F3'))
    tr_ph.append(make_cell('', fill='D9E2F3'))
    tbl.append(tr_ph)
    
    # Phase summary
    all_tids = [t for _, _, ts in subs for t in ts]
    total = sum(T[t]["dur"] for t in all_tids)
    weeks = []
    for t in all_tids:
        for p in wk(T[t]["mulai"], T[t]["selesai"]).split('–'):
            weeks.append(int(p))
    wks = str(min(weeks)) if min(weeks) == max(weeks) else f"{min(weeks)}–{max(weeks)}"
    
    tr_sum = OxmlElement('w:tr')
    tr_sum.append(make_cell('', fill='EAF2FA'))
    tr_sum.append(make_cell('', fill='EAF2FA'))
    tr_sum.append(make_cell('', fill='EAF2FA'))
    tr_sum.append(make_cell('', fill='EAF2FA'))
    tr_sum.append(make_cell(wks, center=True, fill='EAF2FA'))
    tr_sum.append(make_cell(f"{total} hr", center=True, fill='EAF2FA'))
    tr_sum.append(make_cell('', fill='EAF2FA'))
    tbl.append(tr_sum)
    
    for sk, sn, tids in subs:
        if sk:
            tr_sh = OxmlElement('w:tr')
            tr_sh.append(make_cell('', fill='E8F0FE'))
            tr_sh.append(make_cell(sk, bold=True, fill='E8F0FE'))
            tr_sh.append(make_cell(sn, bold=True, fill='E8F0FE'))
            tr_sh.append(make_cell('', fill='E8F0FE'))
            tr_sh.append(make_cell('', fill='E8F0FE'))
            tr_sh.append(make_cell('', fill='E8F0FE'))
            tr_sh.append(make_cell('', fill='E8F0FE'))
            tbl.append(tr_sh)
        
        for tid in tids:
            t = T[tid]
            seq += 1
            fill = 'F2F7FC' if row_idx % 2 == 1 else None
            row_idx += 1
            w = wk(t["mulai"], t["selesai"])
            tr = OxmlElement('w:tr')
            tr.append(make_cell(str(seq), center=True, fill=fill, size=16))
            tr.append(make_cell(t["kode"], center=True, fill=fill, size=16))
            tr.append(make_cell(t["nama"], fill=fill, size=16))
            tr.append(make_cell(t["pic"], center=True, fill=fill, size=16))
            tr.append(make_cell(w, center=True, fill=fill, size=16))
            tr.append(make_cell(f"{t['dur']} hr", center=True, fill=fill, size=16))
            tr.append(make_cell(t["ket"], fill=fill, size=16))
            tbl.append(tr)

insert_after.addnext(tbl)
insert_after = tbl

# Note
sp2 = OxmlElement('w:p')
insert_after.addnext(sp2)
insert_after = sp2
note = make_para("Catatan: Fase berjalan sekuensial, namun task dengan PIC berbeda dalam satu fase berjalan paralel. Weekend (Sabtu-Minggu) tidak dihitung sebagai hari kerja.", italic=True, size=18)
insert_after.addnext(note)

# Save
doc.save(PATH)
print(f"✅ Done! Narrative + 1 table ({seq} tasks) written to DOCX")
