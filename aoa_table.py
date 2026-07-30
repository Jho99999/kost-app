"""
Inject AOA tables (Event + Aktivitas) into SPMP DOCX after the AOA narrative.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOCX_PATH = r"E:\UNIKOM\semester 6\Manajemen Proyek Perangkat Lunak\TUBES\SPMP - Sistem Informasi Manajemen Kost.docx"

doc = Document(DOCX_PATH)

def make_cell(table, row, col, text, bold=False, size=9, color=None, align='center'):
    cell = table.cell(row, col)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {
        'center': 1, 'left': 0, 'right': 2
    }.get(align, 1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set cell vertical alignment middle
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(tcVAlign)
    # Cell borders
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def shade_cell(table, row, col, color="D9E2F3"):
    cell = table.cell(row, col)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)

# ===== FIND INSERTION POINT =====
# We want to add tables after the AOA narrative and before "D.4 Gantt Chart"
# Find the paragraph containing "D.4 Gantt Chart"
insert_idx = None
for i, para in enumerate(doc.paragraphs):
    if "D.4 Gantt Chart" in para.text:
        insert_idx = i
        break

if insert_idx is None:
    # Fallback: look for "Akan dilengkapi kemudian"
    for i, para in enumerate(doc.paragraphs):
        if "Akan dilengkapi kemudian" in para.text and "Gantt" in para.text:
            insert_idx = i
            break

if insert_idx is None:
    print("⚠️  Could not find D.4 Gantt Chart insertion point. Appending at end.")
    insert_idx = len(doc.paragraphs)

# We'll insert BEFORE the D.4 heading
# The paragraph at insert_idx is "D.4 Gantt Chart"
# We insert a new paragraph before it, then the tables

# To insert before a paragraph, we need to work with the XML element tree
body = doc.element.body
d4_para = doc.paragraphs[insert_idx]._element

# ── 1. TABLE AOA EVENT (Forward/Backward Pass) ──

# Insert heading before D.4
heading = parse_xml(
    f'<w:p {nsdecls("w")}>'
    f'<w:pPr><w:pStyle w:val="Heading3"/></w:pPr>'
    f'<w:r><w:t>Tabel Forward/Backward Pass Event AOA</w:t></w:r>'
    f'</w:p>'
)
body.insert(list(body).index(d4_para), heading)

# Event data
events_data = [
    ("1", "—", "A1, A4", "0", "0", "0", "Ya"),
    ("2", "A1", "A2", "3", "3", "0", "Ya"),
    ("3", "A2", "A3", "6", "6", "0", "Ya"),
    ("4", "A3, A4", "B1, B2", "9", "9", "0", "Ya"),
    ("5", "B1, D1", "C1, C3", "24", "24", "0", "Ya"),
    ("6", "B2", "D1", "13", "24", "11", "Tidak"),
    ("7", "C1", "C2", "46", "46", "0", "Ya"),
    ("8", "C2, D2", "D3", "69", "69", "0", "Ya"),
    ("9", "C3", "D2", "36", "69", "33", "Tidak"),
    ("10", "D3", "E1, E2", "81", "81", "0", "Ya"),
    ("11", "E1", "D4", "83", "84", "1", "Tidak"),
    ("12", "E2, D4", "—", "84", "84", "0", "Ya"),
]

cols = 7
rows = len(events_data) + 1  # header + data
tbl1 = doc.add_table(rows=rows, cols=cols)
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
headers = ["Event", "Aktivitas\nMasuk", "Aktivitas\nKeluar", "ES\n(Hari ke-)", "LF\n(Hari ke-)", "Float\n(Hari)", "Critical?"]
for c, h in enumerate(headers):
    make_cell(tbl1, 0, c, h, bold=True, size=8, color=(0, 0, 0))
    shade_cell(tbl1, 0, c, "D9E2F3")

# Data
for r, (ev, masuk, keluar, es, lf, fl, crit) in enumerate(events_data, 1):
    make_cell(tbl1, r, 0, ev, size=8, align='center')
    make_cell(tbl1, r, 1, masuk, size=7, align='center')
    make_cell(tbl1, r, 2, keluar, size=7, align='center')
    make_cell(tbl1, r, 3, es, size=8, align='center')
    make_cell(tbl1, r, 4, lf, size=8, align='center')
    make_cell(tbl1, r, 5, fl, size=8, align='center')
    make_cell(tbl1, r, 6, crit, size=8, align='center')
    if crit == "Ya":
        shade_cell(tbl1, r, 0, "E2EFDA")  # light green for critical
        shade_cell(tbl1, r, 6, "E2EFDA")
    else:
        shade_cell(tbl1, r, 0, "FCE4EC")  # light pink for non-critical
        shade_cell(tbl1, r, 6, "FCE4EC")

# Set column widths
for row in tbl1.rows:
    row.cells[0].width = Cm(1.2)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(2.5)
    row.cells[3].width = Cm(2.0)
    row.cells[4].width = Cm(2.0)
    row.cells[5].width = Cm(1.5)
    row.cells[6].width = Cm(1.5)

# Add spacing paragraph after table
sp1 = parse_xml(
    f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="120"/></w:pPr></w:p>'
)
tbl1._element.addnext(sp1)

# Move table before D.4
body.remove(tbl1._element)
body.insert(list(body).index(d4_para), tbl1._element)
body.remove(sp1)
body.insert(list(body).index(d4_para), sp1)

# ── 2. TABLE AKTIVITAS AOA ──
act_heading = parse_xml(
    f'<w:p {nsdecls("w")}>'
    f'<w:pPr><w:pStyle w:val="Heading3"/></w:pPr>'
    f'<w:r><w:t>Tabel Aktivitas AOA</w:t></w:r>'
    f'</w:p>'
)
body.insert(list(body).index(d4_para), act_heading)

activities = [
    ("1", "A1", "Inisiasi & Identifikasi", "3", "1", "2", "Ya"),
    ("2", "A2", "Fitur & Perencanaan", "3", "2", "3", "Ya"),
    ("3", "A3", "Risiko & SPMP", "3", "3", "4", "Ya"),
    ("4", "A4", "AI Cari Referensi", "1", "1", "4", "Tidak"),
    ("5", "B1", "Desain UI/UX (Admin & Pengguna)", "15", "4", "5", "Ya"),
    ("6", "B2", "Setup Lingkungan (Laravel + Hosting)", "4", "4", "6", "Tidak"),
    ("7", "D1", "Dummy (B2→B1)", "0", "6", "5", "Tidak"),
    ("8", "C1", "Frontend", "22", "5", "7", "Ya"),
    ("9", "C2", "Backend", "23", "7", "8", "Ya"),
    ("10", "C3", "Database", "12", "5", "9", "Tidak"),
    ("11", "D2", "Dummy (C3→C2)", "0", "9", "8", "Tidak"),
    ("12", "D3", "Pengujian & UAT", "12", "8", "10", "Ya"),
    ("13", "E1", "Deployment & SSL", "2", "10", "11", "Tidak"),
    ("14", "E2", "Pelatihan & Dokumentasi", "3", "10", "12", "Ya"),
    ("15", "D4", "Dummy (E1→E2)", "0", "11", "12", "Tidak"),
]

cols2 = 8
rows2 = len(activities) + 1
tbl2 = doc.add_table(rows=rows2, cols=cols2)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ["No", "Aktivitas", "Nama Aktivitas", "Durasi\n(Hari)", "Dari\nEvent", "Ke\nEvent", "Jalur\nKritis?", "ES→EF\nEvent"]
for c, h in enumerate(headers2):
    make_cell(tbl2, 0, c, h, bold=True, size=8, color=(0,0,0))
    shade_cell(tbl2, 0, c, "D9E2F3")

# ES for events (from forward pass)
es_ev = {"1":0, "2":3, "3":6, "4":9, "5":24, "6":13, "7":46, "8":69, "9":36, "10":81, "11":83, "12":84}

for r, (no, kode, nama, dur, fr, to, crit) in enumerate(activities, 1):
    es_from = es_ev[fr]
    es_to = es_ev[to]
    es_ef_str = f"{es_from}→{es_to}"
    
    make_cell(tbl2, r, 0, no, size=8, align='center')
    make_cell(tbl2, r, 1, kode, size=8, align='center')
    make_cell(tbl2, r, 2, nama, size=7, align='left')
    make_cell(tbl2, r, 3, dur, size=8, align='center')
    make_cell(tbl2, r, 4, fr, size=8, align='center')
    make_cell(tbl2, r, 5, to, size=8, align='center')
    make_cell(tbl2, r, 6, crit, size=8, align='center')
    make_cell(tbl2, r, 7, es_ef_str, size=7, align='center')
    
    if crit == "Ya":
        shade_cell(tbl2, r, 0, "E2EFDA")
        shade_cell(tbl2, r, 6, "E2EFDA")
    else:
        shade_cell(tbl2, r, 0, "FCE4EC")
        shade_cell(tbl2, r, 6, "FCE4EC")

# Spacing before D.4
sp2 = parse_xml(
    f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="120"/></w:pPr></w:p>'
)
tbl2._element.addnext(sp2)

# Move table before D.4
body.remove(tbl2._element)
body.insert(list(body).index(d4_para), tbl2._element)
body.remove(sp2)
body.insert(list(body).index(d4_para), sp2)

doc.save(DOCX_PATH)
print(f"✅ Tabel AOA berhasil diinject ke DOCX")
print(f"   - Tabel Event AOA: {rows} baris (12 event)")
print(f"   - Tabel Aktivitas AOA: {rows2} baris (15 aktivitas + dummy)")
