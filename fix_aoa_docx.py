"""
Update DOCX: 
1. Fix narrative (event 1-12 → A-L)
2. Remove old AOA tables, add clean table: ID, Deskripsi Aktivitas, Durasi, Mulai, Selesai
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

DOCX_PATH = r"E:\UNIKOM\semester 6\Manajemen Proyek Perangkat Lunak\TUBES\SPMP - Sistem Informasi Manajemen Kost.docx"

doc = Document(DOCX_PATH)

def make_cell(table, row, col, text, bold=False, size=9, color=None, align='center'):
    cell = table.cell(row, col)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': 1, 'left': 0, 'right': 2}.get(align, 1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(tcVAlign)
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def shade_cell(table, row, col, color="D9E2F3"):
    tc = table.cell(row, col)._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)

# ===== 1. FIX NARRATIVE: replace event numbers with letters =====
num_to_letter = {"1":"A","2":"B","3":"C","4":"D","5":"E","6":"F",
                 "7":"G","8":"H","9":"I","10":"J","11":"K","12":"L"}

# Find the AOA narrative paragraphs and fix them
for para in doc.paragraphs:
    text = para.text
    
    # Fix the critical path sentence: "1 → 2 → 3 → 4 → 5 → 7 → 8 → 10 → 12"
    if "jalur kritis" in text.lower() and "1 → 2" in text:
        new_text = text
        for n, l in num_to_letter.items():
            # Only replace standalone numbers with arrow context
            new_text = new_text.replace(f" {n} →", f" {l} →")
            new_text = new_text.replace(f"→ {n} ", f"→ {l} ")
            new_text = new_text.replace(f"→ {n}", f"→ {l}")  # end of line
        # Fix edge cases
        for r in para.runs:
            if "1 → 2" in r.text or "jalur kritis" in r.text.lower():
                r.text = new_text
                break
        continue
    
    # Fix "event: X → Y" patterns
    if re.search(r'event:?\s*\d+\s*→', text.lower()):
        new_text = text
        for n, l in num_to_letter.items():
            new_text = new_text.replace(f" {n} →", f" {l} →")
            new_text = new_text.replace(f"→ {n}", f"→ {l}")
            # Also fix standalone numbers like "event 1" → "event A"
            new_text = re.sub(rf'\bevent {n}\b', f'event {l}', new_text)
        for r in para.runs:
            if any(d in r.text for d in ["1 →", "2 →", "event"]):
                r.text = new_text
                break

# Also fix in tables (the old AOA table with numbers)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                txt = para.text.strip()
                # If cell contains only a number that matches our events
                if txt in num_to_letter and txt.isdigit():
                    for run in para.runs:
                        run.text = num_to_letter[txt]

# ===== 2. FIND AND REMOVE OLD AOA TABLES =====
# Tables to remove:
# - Table with header containing "Aktivitas" and "Dur" and "Dari" (old AOA table)
# - Table with header "Event" and "Aktivitas Masuk" (Event forward/backward)
# - Table with header "Aktivitas" and "Nama Aktivitas" and "Dari Event" (Activity AOA table)

body = doc.element.body
tables_to_remove = []

for i, table in enumerate(doc.tables):
    first_cell_text = ""
    if table.rows and table.rows[0].cells:
        first_cell_text = " ".join(c.text.strip() for c in table.rows[0].cells[:3])
    
    # Identify AOA tables
    is_aoa_event_table = ("Event" in first_cell_text and "Aktivitas" in first_cell_text and "Masuk" in first_cell_text)
    is_aoa_activity_table = ("Aktivitas" in first_cell_text and "Nama Aktivitas" in first_cell_text)
    is_old_aoa_table = ("Aktivitas" in first_cell_text and "Nama" in first_cell_text and "Dur" in first_cell_text)

    if is_aoa_event_table or is_aoa_activity_table or is_old_aoa_table:
        # Verify it's not the AON table by checking for "Event" in first column data
        if table.rows and len(table.rows) > 2:
            second_col = table.rows[1].cells[0].text.strip() if table.rows[1].cells else ""
            if second_col in num_to_letter or second_col == "1" or second_col == "A1":
                tables_to_remove.append(table._element)
                print(f"  → Removing table {i}: header=[{first_cell_text[:50]}]")

# Also remove adjacent headings (Tabel Forward/Backward Pass Event AOA, Tabel Aktivitas AOA)
for para in list(doc.paragraphs):
    if para.text.strip() in [
        "Tabel Forward/Backward Pass Event AOA",
        "Tabel Aktivitas AOA"
    ]:
        # Find and remove the paragraph element
        p_elem = para._element
        try:
            body.remove(p_elem)
            print(f"  → Removing heading: {para.text.strip()}")
        except:
            pass

# Remove the identified tables
for tbl_elem in tables_to_remove:
    try:
        body.remove(tbl_elem)
        print(f"  → Table element removed")
    except:
        pass

# ===== 3. ADD NEW CLEAN TABLE: ID, Deskripsi Aktivitas, Durasi, Mulai, Selesai =====

# Find insertion point: before "D.4 Gantt Chart"
insert_idx = None
for i, para in enumerate(doc.paragraphs):
    if "D.4 Gantt Chart" in para.text:
        insert_idx = i
        break

if insert_idx is None:
    print("⚠️  D.4 not found, appending")
    insert_idx = len(doc.paragraphs)

d4_para = doc.paragraphs[insert_idx]._element

# Add heading
heading = parse_xml(
    f'<w:p {nsdecls("w")}>'
    f'<w:pPr><w:pStyle w:val="Heading3"/></w:pPr>'
    f'<w:r><w:t>Tabel Perhitungan AOA</w:t></w:r>'
    f'</w:p>'
)
body.insert(list(body).index(d4_para), heading)

# Activity data with ES/EF
activities = [
    ("A1", "Inisiasi & Identifikasi", "3", "0 (H-0)", "3 (H-3)", "Ya"),
    ("A2", "Fitur & Perencanaan", "3", "3 (H-3)", "6 (H-6)", "Ya"),
    ("A3", "Risiko & SPMP", "3", "6 (H-6)", "9 (H-9)", "Ya"),
    ("A4", "AI Cari Referensi", "1", "0 (H-0)", "1 (H-1)", "Tidak"),
    ("B1", "Desain UI/UX (Admin & Pengguna)", "15", "9 (H-9)", "24 (H-24)", "Ya"),
    ("B2", "Setup Lingkungan (Laravel + Hosting)", "4", "9 (H-9)", "13 (H-13)", "Tidak"),
    ("D1", "Dummy (B2→B1)", "0", "13 (H-13)", "24 (H-24)", "Tidak"),
    ("C1", "Frontend", "22", "24 (H-24)", "46 (H-46)", "Ya"),
    ("C2", "Backend", "23", "46 (H-46)", "69 (H-69)", "Ya"),
    ("C3", "Database", "12", "24 (H-24)", "36 (H-36)", "Tidak"),
    ("D2", "Dummy (C3→C2)", "0", "36 (H-36)", "69 (H-69)", "Tidak"),
    ("D3", "Pengujian & UAT", "12", "69 (H-69)", "81 (H-81)", "Ya"),
    ("E1", "Deployment & SSL", "2", "81 (H-81)", "83 (H-83)", "Tidak"),
    ("E2", "Pelatihan & Dokumentasi", "3", "81 (H-81)", "84 (H-84)", "Ya"),
    ("D4", "Dummy (E1→E2)", "0", "83 (H-83)", "84 (H-84)", "Tidak"),
]

cols = 6
rows = len(activities) + 1
tbl = doc.add_table(rows=rows, cols=cols)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["ID", "Deskripsi Aktivitas", "Durasi (Hari)", "Mulai (ES)", "Selesai (EF)", "Jalur Kritis?"]
col_widths = [1.2, 5.0, 1.8, 1.8, 1.8, 1.5]  # in cm approx

for c, h in enumerate(headers):
    make_cell(tbl, 0, c, h, bold=True, size=8, color=(0,0,0))
    shade_cell(tbl, 0, c, "D9E2F3")

for r, (id_, desc, dur, mulai, selesai, crit) in enumerate(activities, 1):
    make_cell(tbl, r, 0, id_, size=8, align='center')
    make_cell(tbl, r, 1, desc, size=7, align='left')
    make_cell(tbl, r, 2, dur, size=8, align='center')
    make_cell(tbl, r, 3, mulai, size=7, align='center')
    make_cell(tbl, r, 4, selesai, size=7, align='center')
    make_cell(tbl, r, 5, crit, size=8, align='center')
    if crit == "Ya":
        shade_cell(tbl, r, 0, "E2EFDA")
        shade_cell(tbl, r, 5, "E2EFDA")
    else:
        shade_cell(tbl, r, 0, "FCE4EC")
        shade_cell(tbl, r, 5, "FCE4EC")

# Move table before D.4
sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="120"/></w:pPr></w:p>')
tbl._element.addnext(sp)
body.remove(tbl._element)
body.insert(list(body).index(d4_para), tbl._element)
body.remove(sp)
body.insert(list(body).index(d4_para), sp)

doc.save(DOCX_PATH)
print(f"\n✅ DOCX updated!")
print(f"   - Event labels: 1-12 → A-L")
print(f"   - Old AOA tables removed")
print(f"   - New table: {rows-1} activities with ES/EF")
