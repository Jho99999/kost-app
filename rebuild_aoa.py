"""
REBUILD AOA section from scratch.
Remove everything between "Diagram AOA" and "D.4 Gantt Chart" and replace with clean content.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

DOCX_PATH = r"E:\UNIKOM\semester 6\Manajemen Proyek Perangkat Lunak\TUBES\SPMP - Sistem Informasi Manajemen Kost.docx"

doc = Document(DOCX_PATH)
body = doc.element.body

# ===== FIND THE RANGE TO REPLACE =====
start_para = None
end_para = None

for i, para in enumerate(doc.paragraphs):
    if "Diagram AOA" in para.text and "Activity on Arrow" in para.text:
        start_para = para._element
    if "D.4 Gantt Chart" in para.text:
        end_para = para._element
        break

if start_para is None or end_para is None:
    print("❌ Could not find AOA section boundaries")
    print(f"  start_para={'found' if start_para else 'MISSING'}")
    print(f"  end_para={'found' if end_para else 'MISSING'}")
    exit(1)

# ===== REMOVE EVERYTHING BETWEEN start_para AND end_para =====
# Also track tables to remove (any table between these paragraphs)
children = list(body)
in_range = False
elements_to_remove = []
tables_found = 0

for child in children:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    
    if child is start_para:
        in_range = True
        # Keep start_para itself (we'll modify it)
        continue
    
    if child is end_para:
        break  # stop before D.4
    
    if in_range:
        if tag in ['p', 'tbl']:
            elements_to_remove.append(child)
            if tag == 'tbl':
                tables_found += 1

print(f"Found AOA section: {len(elements_to_remove)} elements ({tables_found} tables)")
for e in elements_to_remove:
    try:
        body.remove(e)
    except:
        pass

# ===== REBUILD: Fix the start paragraph text (event numbers A-L) =====
# Replace the start paragraph text: "Diagram AOA proyek ini memiliki 12 event node..."
for para in doc.paragraphs:
    if para._element is start_para:
        # Actually, find the next paragraph after "Diagram AOA (Activity on Arrow)" heading
        break

# Find the paragraph that describes the 12 events and fix event refs
for para in doc.paragraphs:
    if "12 event node" in para.text:
        for run in para.runs:
            run.text = run.text.replace(
                "event: 1 → 2 → 3 → 4 → 5 → 7 → 8 → 10 → 12",
                "event: A → B → C → D → E → G → H → J → L"
            )
            run.text = run.text.replace(
                "1 → 2 → 3 → 4 → 5 → 7 → 8 → 10 → 12",
                "A → B → C → D → E → G → H → J → L"
            )
        break

# Also fix the critical path description if it exists
for para in doc.paragraphs:
    if "jalur kritis (critical path)" in para.text.lower():
        for run in para.runs:
            run.text = run.text.replace(
                "event: 1 → 2 → 3 → 4 → 5 → 7 → 8 → 10 → 12",
                "event: A → B → C → D → E → G → H → J → L"
            )
        break

# Fix the old "12 event node" description text
for para in doc.paragraphs:
    if "12 event" in para.text:
        original = para.text
        # Replace "12 aktivitas (panah solid)" and event numbers in context
        for run in para.runs:
            t = run.text
            t = t.replace("12 event node", "12 event (A–L)")
            t = t.replace("12 event", "12 event (A–L)")
            run.text = t
        break

# ===== INSERT NEW TABLE =====
# Helper
def make_cell(table, row, col, text, bold=False, size=9, color=None, align='center'):
    cell = table.cell(row, col)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'center': 1, 'left': 0, 'right': 2}.get(align, 1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(tcVAlign)
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def shade_cell(table, row, col, color="D9E2F3"):
    tc = table.cell(row, col)._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)

# Insert heading
heading_elem = parse_xml(
    f'<w:p {nsdecls("w")}>'
    f'<w:pPr><w:pStyle w:val="Heading3"/></w:pPr>'
    f'<w:r><w:t>Tabel Perhitungan AOA</w:t></w:r>'
    f'</w:p>'
)

# Insert before end_para (D.4)
end_idx = list(body).index(end_para)
body.insert(end_idx, heading_elem)

# Activity data
activities = [
    ("A1", "Inisiasi & Identifikasi", "3", "0 (H-0)", "3 (H-3)", "Ya"),
    ("A2", "Fitur & Perencanaan", "3", "3 (H-3)", "6 (H-6)", "Ya"),
    ("A3", "Risiko & SPMP", "3", "6 (H-6)", "9 (H-9)", "Ya"),
    ("A4", "AI Cari Referensi", "1", "0 (H-0)", "1 (H-1)", "Tidak"),
    ("B1", "Desain UI/UX (Admin & Pengguna)", "15", "9 (H-9)", "24 (H-24)", "Ya"),
    ("B2", "Setup Lingkungan (Laravel + Hosting)", "4", "9 (H-9)", "13 (H-13)", "Tidak"),
    ("D1", "Dummy (B2→B1)", "0", "13 (H-13)", "24 (H-24)", "Tidak"),
    ("C1", "Frontend", "22", "24 (H-24)", "46 (H-46)", "Ya"),
    ("C2", "Backend", "23", "46 (H-46)", "69 (H-69)", "Ya"),
    ("C3", "Database", "12", "24 (H-24)", "36 (H-36)", "Tidak"),
    ("D2", "Dummy (C3→C2)", "0", "36 (H-36)", "69 (H-69)", "Tidak"),
    ("D3", "Pengujian & UAT", "12", "69 (H-69)", "81 (H-81)", "Ya"),
    ("E1", "Deployment & SSL", "2", "81 (H-81)", "83 (H-83)", "Tidak"),
    ("E2", "Pelatihan & Dokumentasi", "3", "81 (H-81)", "84 (H-84)", "Ya"),
    ("D4", "Dummy (E1→E2)", "0", "83 (H-83)", "84 (H-84)", "Tidak"),
]

cols = 6
rows = len(activities) + 1
tbl = doc.add_table(rows=rows, cols=cols)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["ID", "Deskripsi Aktivitas", "Durasi", "Mulai (ES)", "Selesai (EF)", "Kritis?"]
for c, h in enumerate(headers):
    make_cell(tbl, 0, c, h, bold=True, size=8, color=(0,0,0))
    shade_cell(tbl, 0, c, "D9E2F3")

for r, (id_, desc, dur, mulai, selesai, crit) in enumerate(activities, 1):
    make_cell(tbl, r, 0, id_, size=8, align='center')
    make_cell(tbl, r, 1, desc, size=7, align='left')
    make_cell(tbl, r, 2, dur, size=8, align='center')
    make_cell(tbl, r, 3, mulai, size=7, align='center')
    make_cell(tbl, r, 4, selesai, size=7, align='center')
    make_cell(tbl, r, 5, crit, size=8, align='center')
    if crit == "Ya":
        shade_cell(tbl, r, 0, "E2EFDA")
        shade_cell(tbl, r, 5, "E2EFDA")
    else:
        shade_cell(tbl, r, 0, "FCE4EC")
        shade_cell(tbl, r, 5, "FCE4EC")

# Move table before D.4 using addnext on end_para
# Build: spacer, table, heading (in reverse since addnext inserts after)
sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="120"/></w:pPr></w:p>')

# Remove table from wherever doc.add_table put it (usually at end of body)
body.remove(tbl._element)

# Insert in order: heading → table → spacer → [D.4]
# addnext inserts RIGHT AFTER the element, so we add in reverse
end_para.addnext(sp)               # after sp comes D.4
end_para.addnext(tbl._element)      # after tbl comes sp
end_para.addnext(heading_elem)      # after heading comes tbl

doc.save(DOCX_PATH)
print(f"✅ DOCX AOA section rebuilt!")
print(f"   - Removed old tables + damaged data")
print(f"   - Narrative uses event A–L")
print(f"   - New table: {rows-1} rows (ID, Deskripsi, Durasi, Mulai, Selesai, Kritis?)")
